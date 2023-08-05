from io import BytesIO
import docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml import ns, parse_xml
from docx.shared import Pt
from docx.table import _Cell
from docx2pdf import convert
from datetime import datetime
from num2words import num2words
import pythoncom
import pathlib
from pathlib import Path
import tempfile
import calendar
import locale

locale.setlocale(locale.LC_ALL, 'uk_UA.UTF-8')

def month_number_to_name(month_number):
    month_number = int(month_number)
    return calendar.month_name[month_number]
def replace_text(doc, old_text, new_text):
    for paragraph in doc.paragraphs:
        new_runs = []
        for run in paragraph.runs:
            if old_text in run.text:
                new_run = parse_xml(run._element.xml)
                new_run.text = new_run.text.replace(old_text, new_text)
                new_runs.append(new_run)
            else:
                new_runs.append(run._element)
        paragraph._element.clear_content()
        for new_run in new_runs:
            paragraph._element.append(new_run)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    new_runs = []
                    for run in paragraph.runs:
                        if old_text in run.text:
                            new_run = parse_xml(run._element.xml)
                            new_run.text = new_run.text.replace(old_text, new_text)
                            new_runs.append(new_run)
                        else:
                            new_runs.append(run._element)
                    paragraph._element.clear_content()
                    for new_run in new_runs:
                        paragraph._element.append(new_run)


def make_doc(name_templates, name_company, number_doc, adress, edrpou, ipn, iban, tel, email, other_data, director, money):
    path_to_doc = f'static/docs_template/{name_templates}.docx'
    doc = Document(path_to_doc)
    replace_text(doc, 'КОМПАНІЯ', name_company)
    replace_text(doc, 'НОМЕР_ДОГОВОРУ', number_doc)
    replace_text(doc, 'АДРЕСАГЕН', adress)
    replace_text(doc, 'ЄДРПОУГЕН', str(edrpou))
    replace_text(doc, 'ІПНГЕН', ipn)
    replace_text(doc, 'ІБАНГЕН', iban)
    replace_text(doc, 'ТЕЛЕФОН', tel)
    replace_text(doc, 'ЕПОШТА', email)
    replace_text(doc, 'ІНША_ІНФОРМАЦІЯ',other_data)
    replace_text(doc, 'ПІДПИСАНТ', director)
    if money != '':
        replace_text(doc, 'СУММА', money)
        replace_text(doc, 'ПРОПИСОМ', num2words(int(money), lang='uk'))
    replace_text(doc, 'ДАТА', f'{datetime.today().day}.{datetime.today().month if datetime.today().month < 10 else str(datetime.today().month)}.{datetime.today().year}')
    replace_text(doc, 'МИНУЛИЙМІСЯЦЬПРОПИС', month_number_to_name(
        f'{datetime.today().month - 1:02}' if datetime.today().month - 1 < 10 else str(datetime.today().month - 1)))
    replace_text(doc, 'МИНУЛИЙМІСЯЦЬ', f'{datetime.today().month - 1:02}' if datetime.today().month - 1 < 10 else str(
        datetime.today().month - 1))
    replace_text(doc, 'МІСЯЦЬПРОПИС', month_number_to_name(
        f'{datetime.today().month:02}' if datetime.today().month < 10 else str(datetime.today().month)))
    replace_text(doc, 'МІСЯЦЬ', f'{datetime.today().month:02}' if datetime.today().month < 10 else str(datetime.today().month))
    replace_text(doc, 'ПОВНИЙРІК', f'{datetime.today().year}')
    replace_text(doc, 'РІК', f'{datetime.today().year}'[2:])
    doc.save(path_to_doc)


def doc_to_pdf(name_file):
    pythoncom.CoInitialize()
    convert(f"static/docs_template/{name_file}.docx", f"static/{name_file}.pdf")


def doc_to_pdf_new(name):
    pythoncom.CoInitialize()
    convert(f'{name}.docx', f'{name}.pdf')
    #convert(f"static/docs_template/{name}.docx", f"static/{name}.pdf")

def all_files(): #просматриваем все файлы в директории
    files =[]
    currentPattern = '*.pdf'
    currentDirectory = pathlib.Path('./static')
    for currentFile in currentDirectory.glob(currentPattern):
        files.append(str(currentFile.name).replace('.pdf', ''))
    if 'chenge' in files:
        files.remove('chenge')
    files.remove('Приклад_1')
    files.remove('Приклад_2')
    return files

def delete_files(file_name):
    name_docx = file_name+'.docx'
    # Удаляем docx шаблон
    # Указываем путь к директории
    directory = pathlib.Path('static/docs_template')
    file_path = directory / name_docx

    # Проверяем, существует ли файл
    if file_path.is_file():
        # Удаляем файл
        file_path.unlink()

    # Удаляем PDF шаблон
    # Указываем путь к директории
    name_docx = file_name+'.pdf'
    directory = pathlib.Path('static')
    file_path = directory / name_docx

    # Проверяем, существует ли файл
    if file_path.is_file():
        # Удаляем файл
        file_path.unlink()

def delete_generation_files(id):
    directory = Path(__file__).resolve().parent
    docx_file = directory / f'{id}.docx'
    pdf_file = directory / f'{id}.pdf'

    if docx_file.exists():
        docx_file.unlink()

    if pdf_file.exists():
        pdf_file.unlink()

def make_doc_standart(name_templates, name_company, number_doc, company_details, money):
    path_to_doc = f'static/docs_template/{name_templates}.docx'
    doc = Document(path_to_doc)
    replace_text(doc, '**КОМПАНІЯ**', name_company)
    replace_text(doc, '**НОМЕР_ДОГОВОРУ**', number_doc)
    replace_text(doc, '**РЕКВІЗИТИ_КОМПАНІЇ**', company_details)
    if money != '':
        replace_text(doc, '**СУММА**', money)
        replace_text(doc, '**ПРОПИСОМ**', num2words(int(money), lang='uk'))
    replace_text(doc, '**ДАТА**', f'{datetime.today().day}.{datetime.today().month}.{datetime.today().year}')
    replace_text(doc, '**МІСЯЦЬ**', f'{datetime.today().month}')
    replace_text(doc, '**РІК**', f'{datetime.today().year}'[2:])
    doc.save('static/docs_template/standart_template.docx')

def doc_to_pdf_standart():
    pythoncom.CoInitialize()
    convert('static/docs_template/standart_template.docx', f'static/standart_template.pdf')