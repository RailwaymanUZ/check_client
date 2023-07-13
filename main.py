from docx import Document
from docx.enum.text import WD_UNDERLINE
from lxml import etree
from num2words import num2words
from datetime import datetime

from docx import Document
from docx.enum.text import WD_UNDERLINE
from docx.shared import Pt

from docx import Document
from docx.oxml import ns, parse_xml
from docx.shared import Pt


def replace_text(doc, old_text, new_text):
    for paragraph in doc.paragraphs:
        new_runs = []
        for run in paragraph.runs:
            if old_text in run.text:
                new_run = parse_xml(run._element.xml)
                new_run.text = new_run.text.replace(old_text, new_text)
                new_runs.append(new_run)
            else:
                new_runs.append(run._element)
        paragraph._element.clear_content()
        for new_run in new_runs:
            paragraph._element.append(new_run)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    new_runs = []
                    for run in paragraph.runs:
                        if old_text in run.text:
                            new_run = parse_xml(run._element.xml)
                            new_run.text = new_run.text.replace(old_text, new_text)
                            new_runs.append(new_run)
                        else:
                            new_runs.append(run._element)
                    paragraph._element.clear_content()
                    for new_run in new_runs:
                        paragraph._element.append(new_run)

# Пример использования
def make_doc(name_templates, name_company, number_doc, company_details, money):
    path_to_doc = f'static/docs_template/{name_templates}.docx'
    doc = Document(path_to_doc)
    replace_text(doc, 'КОМПАНІЯ', name_company)
    replace_text(doc, 'НОМЕР_ДОГОВОРУ', number_doc)
    replace_text(doc, 'РЕКВІЗИТИ_КОМПАНІЇ', company_details)
    replace_text(doc, 'СУММА', money)
    replace_text(doc, 'ПРОПИСОМ', num2words(int(money), lang='uk'))
    replace_text(doc, 'ДАТА', f'{datetime.today().day}.{datetime.today().month}.{datetime.today().year}')
    replace_text(doc, 'МІСЯЦЬ', f'{datetime.today().month}')
    replace_text(doc, 'РІК', f'{datetime.today().year}'[2:])
    doc.save('static/docs_template/chenge.docx')









gen_name = 'ТОВ "МАМОНТ-АГРО"'
gen_number = '387550443'
gen_detail = 'ТОВАРИСТВО З ОБМЕЖЕНОЮ ВІДПОВІДАЛЬНІСТЮ "МАМОНТ-АГРО"\n' \
             'Адреса: 02095, Україна, м. Київ, вул. Княжий Затон 21, прим. 566, 576\n' \
             'Код ЄДРПОУ: 41947062\n' \
             'ІПН: 419470626515\n' \
             'IBAN: UA023808050000000026009597181\n' \
             'У банку АТ «Райффайзен Банк Аваль»\n' \
             '\n\n\nДиректор Роман Бундін'
gen_money = '3500'
make_doc('Приклад_1', gen_name, gen_number, gen_detail, gen_money)